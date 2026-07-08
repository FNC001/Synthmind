#!/usr/bin/env python3
from __future__ import annotations

import json
import re
import argparse
import csv
from dataclasses import dataclass
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


@dataclass
class CaseInput:
    key: str
    json_path: Path
    prediction_csv: Path


@dataclass
class ReportCase:
    key: str
    target_formula: str
    structure_summary: str
    structure_source: str
    json_synthesis: str
    json_method: str
    json_precursors: str
    selected: pd.Series
    top_routes: pd.DataFrame


def set_font(run, name: str = "Calibri", east_asia: str = "Songti SC", size: int | None = None) -> None:
    run.font.name = name
    if size is not None:
        run.font.size = Pt(size)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side, val in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                node = tc_mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    tc_mar.append(node)
                node.set(qn("w:w"), str(val))
                node.set(qn("w:type"), "dxa")


def add_para(doc: Document, text: str = "", style: str | None = None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    set_font(r)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_font(r)
    return p


def add_numbered(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_font(r)
    return p


def add_step(doc: Document, num: int, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.22)
    r1 = p.add_run(f"步骤 {num}：")
    r1.bold = True
    set_font(r1)
    r2 = p.add_run(text)
    set_font(r2)
    return p


def cell_text(cell, text: str, bold: bool = False, center: bool = False) -> None:
    cell.text = str(text)
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(0)
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = bold
            set_font(run, size=9)


def configure_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("SynPred pipeline_v3 synthesis-route prediction")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)
    set_font(run, size=9)


def formula_from_json(data: dict) -> str:
    for item in data["results"]:
        if item["task"] in {"method", "precursor"}:
            m = re.search(r'"([^"]+)"', item.get("prompt", ""))
            if m:
                return m.group(1)
    return ""


def get_task(data: dict, task: str) -> dict:
    return next(item for item in data["results"] if item["task"] == task)


def select_route(df: pd.DataFrame) -> pd.Series:
    tmp = df.copy()
    status_rank = {
        "recommended": 0,
        "recommended_with_validation": 1,
        "review_required": 2,
    }
    conf_rank = {"high_confidence": 0, "medium_confidence": 1, "low_confidence": 2}
    tmp["__status"] = tmp.get("condition_distribution_recommendation_status", "").map(status_rank).fillna(3)
    tmp["__conf"] = tmp.get("route_confidence_level", "").map(conf_rank).fillna(3)
    tmp["__qc"] = tmp.get("precursor_qc_status", "").map(status_rank).fillna(3)
    tmp = tmp.sort_values(
        ["__status", "__qc", "__conf", "final_route_rank"],
        ascending=[True, True, True, True],
        kind="mergesort",
    )
    return tmp.iloc[0]


def load_cases(case_inputs: list[CaseInput]) -> list[ReportCase]:
    cases: list[ReportCase] = []
    for item in case_inputs:
        data = json.loads(item.json_path.read_text(encoding="utf-8"))
        df = pd.read_csv(item.prediction_csv)
        formula = formula_from_json(data)
        selected = select_route(df)
        precursor_output = get_task(data, "precursor").get("output", "")
        if formula == "Li4SmP2S8Cl" and precursor_output == "['Th']":
            precursor_output = "raw_output suggests Li2S, P2S5, SmCl3; structured output field is inconsistent"
        cases.append(
            ReportCase(
                key=item.key,
                target_formula=formula,
                structure_summary=data["structure_description"][:280]
                + ("..." if len(data["structure_description"]) > 280 else ""),
                structure_source=data.get("description_source", ""),
                json_synthesis=get_task(data, "synthesis").get("output", ""),
                json_method=get_task(data, "method").get("output", ""),
                json_precursors=precursor_output,
                selected=selected,
                top_routes=df.head(5),
            )
        )
    return cases


def fmt_num(value, digits=1) -> str:
    try:
        if pd.isna(value):
            return "-"
        return f"{float(value):.{digits}f}"
    except Exception:
        return str(value)


def zh_conf(value) -> str:
    return {
        "high_confidence": "高",
        "medium_confidence": "中",
        "low_confidence": "低",
        "high": "高",
        "medium": "中",
        "low": "低",
    }.get(str(value), str(value) if str(value) != "nan" else "-")


def zh_status(value) -> str:
    return {
        "recommended": "推荐",
        "recommended_with_validation": "建议验证",
        "review_required": "需复核",
    }.get(str(value), str(value) if str(value) != "nan" else "-")


def zh_qc(value) -> str:
    return {
        "minor_warning": "轻微警告",
        "major_warning": "主要警告",
        "fail": "未通过",
        "pass": "通过",
    }.get(str(value), str(value) if str(value) != "nan" else "-")


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("三结构合成工艺方案预测报告")
    r.font.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string("0B2545")
    set_font(r, size=22)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("基于远程 SynPred pipeline_v3 单结构推理结果整理")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(85, 85, 85)
    set_font(r)

    add_para(
        doc,
        "本报告采用远程合成路线算法输出的 final_top_routes_with_condition_confidence.csv 作为主依据，"
        "并结合输入 JSON 中的 three_tasks 结果做结构与任务来源说明。所有路线均为模型建议，需经实验安全评估和相纯度验证。",
    )


def add_overview(doc: Document, cases: list[ReportCase]) -> None:
    doc.add_heading("1. 预测总览", level=1)
    table = doc.add_table(rows=len(cases) + 1, cols=6)
    table.style = "Table Grid"
    headers = ["结构", "目标化学式", "推荐前驱体", "温度 / 时间", "气氛", "置信与状态"]
    for i, h in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], "F2F4F7")
        cell_text(table.rows[0].cells[i], h, bold=True, center=True)
    for r_idx, case in enumerate(cases, 1):
        row = table.rows[r_idx].cells
        s = case.selected
        values = [
            case.key,
            case.target_formula,
            s["precursor_set"],
            f"{fmt_num(s['temperature_c'])} °C / {fmt_num(s['time_h'])} h",
            str(s.get("pred_atmosphere", "-")),
            f"{zh_conf(s.get('route_confidence_level', '-'))}; {zh_status(s.get('condition_distribution_recommendation_status', '-'))}",
        ]
        for i, v in enumerate(values):
            cell_text(row[i], v, center=i in {1, 3, 4})
    set_table_geometry(table, [1550, 1250, 2900, 1500, 950, 1210])


def add_case(doc: Document, case: ReportCase, idx: int) -> None:
    doc.add_heading(f"{idx}. {case.key}（{case.target_formula}）", level=1)
    s = case.selected
    add_para(
        doc,
        f"推荐优先路线：{s['precursor_set']}；预测条件为 {fmt_num(s['temperature_c'])} °C、"
        f"{fmt_num(s['time_h'])} h，气氛为 {s.get('pred_atmosphere', '-')}。"
    )

    doc.add_heading("建议工艺步骤", level=2)
    add_step(doc, 1, f"按目标化学计量称取前驱体：{s['precursor_set']}。")
    add_step(doc, 2, "充分研磨混匀；若含易吸湿、挥发或腐蚀性组分，应在受控气氛中处理并选择合适容器。")
    add_step(doc, 3, f"按模型推荐条件升温至约 {fmt_num(s['temperature_c'])} °C，保温约 {fmt_num(s['time_h'])} h。")
    add_step(doc, 4, "冷却后进行 PXRD、Raman/FTIR、元素分析或热分析，确认目标相、残余前驱体和副相。")
    add_step(doc, 5, "若出现副相，优先围绕前驱体组合、气氛、保温时间和二次研磨复烧进行小范围参数扫描。")

    doc.add_heading("算法输出与 QC", level=2)
    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    fields = [
        ("route confidence", f"{zh_conf(s.get('route_confidence_level', '-'))} ({fmt_num(s.get('route_confidence_score'), 3)})"),
        ("condition support", f"{zh_conf(s.get('condition_distribution_confidence_level', '-'))} ({fmt_num(s.get('condition_distribution_support_score'), 3)})"),
        ("recommendation status", zh_status(s.get("condition_distribution_recommendation_status", "-"))),
        ("precursor QC", f"{zh_qc(s.get('precursor_qc_level', '-'))} / {zh_status(s.get('precursor_qc_status', '-'))}"),
        ("precursor QC warnings", s.get("precursor_qc_warnings", "-") if not pd.isna(s.get("precursor_qc_warnings", "")) else "-"),
        ("three_tasks baseline", f"synthesis={case.json_synthesis}; method={case.json_method}; precursor={case.json_precursors}"),
    ]
    for row, (k, v) in zip(table.rows, fields):
        set_cell_shading(row.cells[0], "E8EEF5")
        cell_text(row.cells[0], k, bold=True)
        cell_text(row.cells[1], str(v))
    set_table_geometry(table, [2200, 7160])

    doc.add_heading("Top 5 候选路线", level=2)
    t = doc.add_table(rows=6, cols=6)
    t.style = "Table Grid"
    headers = ["rank", "前驱体", "温度", "时间", "置信", "状态"]
    for i, h in enumerate(headers):
        set_cell_shading(t.rows[0].cells[i], "F2F4F7")
        cell_text(t.rows[0].cells[i], h, bold=True, center=True)
    for r_idx, (_, row_data) in enumerate(case.top_routes.iterrows(), 1):
        vals = [
            int(row_data["final_route_rank"]),
            row_data["precursor_set"],
            f"{fmt_num(row_data['temperature_c'])} °C",
            f"{fmt_num(row_data['time_h'])} h",
            zh_conf(row_data.get("route_confidence_level", "-")),
            zh_status(row_data.get("condition_distribution_recommendation_status", "-")),
        ]
        for i, v in enumerate(vals):
            cell_text(t.rows[r_idx].cells[i], str(v), center=i in {0, 2, 3})
    set_table_geometry(t, [650, 3150, 1150, 1050, 1600, 1760])

    doc.add_heading("结构输入摘要", level=2)
    add_para(doc, f"source={case.structure_source}; {case.structure_summary}")


def add_notes(doc: Document) -> None:
    doc.add_heading("附录：运行口径与限制", level=1)
    add_bullet(doc, "远程流水线按每个结构单独运行，避免批量模式中 precursor_set 全局去重导致样本被误删。")
    add_bullet(doc, "输入 POSCAR 由 JSON 中的 pyxtal 空间群、晶胞和代表性坐标经 pymatgen 对称展开生成；报告保留原目标化学式。")
    add_bullet(doc, "远程任务均在 final_top_routes_with_condition_confidence.csv 生成后，于最后 QC 合并步骤遇到 pandas dtype 写入错误；本报告采用已生成的路线与条件置信度表。")
    add_bullet(doc, "模型输出不是实验 SOP；涉及 SO3、硫化物、卤化物、硝酸盐、铵盐等体系时，应先做安全和兼容性审查。")


def read_manifest(path: Path) -> list[CaseInput]:
    rows: list[CaseInput] = []
    with path.open("r", encoding="utf-8", newline="") as f:
        reader = csv.DictReader(f)
        required = {"key", "json_path", "prediction_csv"}
        missing = required - set(reader.fieldnames or [])
        if missing:
            raise ValueError(f"manifest missing columns: {', '.join(sorted(missing))}")
        for row in reader:
            rows.append(
                CaseInput(
                    key=str(row["key"]).strip(),
                    json_path=Path(str(row["json_path"]).strip()).expanduser(),
                    prediction_csv=Path(str(row["prediction_csv"]).strip()).expanduser(),
                )
            )
    return rows


def discover_inputs(json_dir: Path, prediction_dir: Path) -> list[CaseInput]:
    rows: list[CaseInput] = []
    for json_path in sorted(json_dir.glob("*_three_tasks.json")):
        key = json_path.stem.replace("_three_tasks", "")
        prediction_csv = prediction_dir / f"{key}_final_top_routes_with_condition_confidence.csv"
        rows.append(CaseInput(key=key, json_path=json_path, prediction_csv=prediction_csv))
    if not rows:
        raise FileNotFoundError(f"no *_three_tasks.json files found in {json_dir}")
    return rows


def validate_inputs(case_inputs: list[CaseInput]) -> None:
    if not case_inputs:
        raise ValueError("no cases provided")
    for item in case_inputs:
        if not item.key:
            raise ValueError("case key cannot be empty")
        if not item.json_path.exists():
            raise FileNotFoundError(f"missing JSON for {item.key}: {item.json_path}")
        if not item.prediction_csv.exists():
            raise FileNotFoundError(f"missing prediction CSV for {item.key}: {item.prediction_csv}")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build a DOCX synthesis-route report from three_tasks JSON and SynPred route CSV files.")
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--manifest", type=Path, help="CSV with key,json_path,prediction_csv columns.")
    group.add_argument("--json-dir", type=Path, help="Directory containing *_three_tasks.json files.")
    parser.add_argument(
        "--prediction-dir",
        type=Path,
        help="Directory containing <key>_final_top_routes_with_condition_confidence.csv files. Required with --json-dir.",
    )
    parser.add_argument("--output", type=Path, required=True, help="Output DOCX path.")
    return parser.parse_args()


def build(case_inputs: list[CaseInput], output: Path) -> None:
    cases = load_cases(case_inputs)
    doc = Document()
    configure_doc(doc)
    add_title(doc)
    add_overview(doc, cases)
    for idx, case in enumerate(cases, 2):
        add_case(doc, case, idx)
    add_notes(doc)
    doc.core_properties.title = "三结构合成工艺方案预测报告"
    doc.core_properties.subject = "SynPred pipeline_v3"
    doc.core_properties.author = "Codex"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(output)


if __name__ == "__main__":
    args = parse_args()
    if args.manifest:
        case_inputs = read_manifest(args.manifest)
    else:
        if args.prediction_dir is None:
            raise SystemExit("--prediction-dir is required with --json-dir")
        case_inputs = discover_inputs(args.json_dir, args.prediction_dir)
    validate_inputs(case_inputs)
    build(case_inputs, args.output)
